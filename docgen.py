"""Generates an editable .docx lead report (plus a matching .csv) from
matched leads.

The report is LANDSCAPE US Letter. Column widths are fixed DXA values
summing to exactly 12960 DXA (9in), the usable width on a landscape
Letter page with 1in margins (11in - 2*1in = 9in = 12960 DXA, since
1in = 1440 DXA). Word will not shrink or overflow the table as long as
this sum holds and autofit is disabled.

It was portrait with four columns (Address / Owner / Distress Score /
Sources Hit), which was the reason the report told the reader to "see
case_no" for an unknown address without carrying a case_no column to see.
Every column the pipeline populates now has a home, which needs the extra
width.
"""
import csv
from datetime import datetime
from typing import List

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from match import ADDRESS_SOURCE_HCAD, ADDRESS_SOURCE_UNKNOWN, MatchedLead

USABLE_WIDTH_DXA = 12960

# DXA widths -- must sum to USABLE_WIDTH_DXA
COL_WIDTHS_DXA = {
    "Address": 3200,
    "Owner": 2400,
    "County": 800,
    "Case No": 1600,
    "Market Value": 1100,
    "Score": 800,
    "Sources Hit": 1800,
    "Addr. Source": 1260,
}
assert sum(COL_WIDTHS_DXA.values()) == USABLE_WIDTH_DXA, \
    f"column widths must sum to {USABLE_WIDTH_DXA} DXA"

COLUMNS = list(COL_WIDTHS_DXA)

SOURCE_LABELS = {
    "tax_delinquent": "Tax Delinquent",
    "probate": "Probate",
    "trustee_sale": "Trustee Sale",
    "absentee_owner": "Absentee Owner",
}

ADDRESS_SOURCE_LABELS = {
    ADDRESS_SOURCE_HCAD: "HCAD owner match",
    ADDRESS_SOURCE_UNKNOWN: "Not found",
}

# Full field set for the CSV sidecar. The .docx is the readable report;
# this is the one to open in Excel or feed back through --manual-csv.
CSV_COLUMNS = [
    "address", "owner_name", "county", "case_no", "sources_hit", "distress_score",
    "address_source", "address_confidence", "hcad_matched_owner", "parcel_id",
    "mailing_address", "market_value", "year_built", "building_sqft", "source_urls",
]


def _money(value) -> str:
    return f"${value:,.0f}" if value else ""


def _address_source_label(lead: MatchedLead) -> str:
    label = ADDRESS_SOURCE_LABELS.get(lead.address_source, "County filing")
    if lead.address_source == ADDRESS_SOURCE_HCAD and lead.address_confidence:
        return f"{label} ({lead.address_confidence:.0f}%)"
    return label


def _set_cell_width(cell, dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _disable_autofit(table):
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _row_values(lead: MatchedLead) -> List[str]:
    return [
        lead.address,
        lead.owner_name,
        lead.county,
        lead.case_numbers,
        _money(lead.market_value),
        str(lead.distress_score),
        ", ".join(SOURCE_LABELS.get(s, s) for s in lead.sources_hit),
        _address_source_label(lead),
    ]


def build_document(leads: List[MatchedLead], generated_at: datetime = None) -> Document:
    generated_at = generated_at or datetime.now()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    title = doc.add_heading("Texas Distress-Lead Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    recovered = sum(1 for lead in leads if lead.address_source == ADDRESS_SOURCE_HCAD)
    unknown = sum(1 for lead in leads if lead.address_source == ADDRESS_SOURCE_UNKNOWN)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Generated {generated_at.strftime('%Y-%m-%d %H:%M')} — "
        f"Harris & Nacogdoches County, TX — {len(leads)} matched properties "
        f"({recovered} address{'es' if recovered != 1 else ''} recovered from the "
        f"HCAD roll, {unknown} still unknown)"
    ).italic = True

    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    _disable_autofit(table)

    for i, col in enumerate(table.columns):
        col.width = Inches(COL_WIDTHS_DXA[COLUMNS[i]] / 1440)

    header_cells = table.rows[0].cells
    for i, name in enumerate(COLUMNS):
        header_cells[i].text = name
        for p in header_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        _set_cell_width(header_cells[i], COL_WIDTHS_DXA[name])

    for lead in leads:
        row = table.add_row().cells
        for i, val in enumerate(_row_values(lead)):
            row[i].text = val
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
            _set_cell_width(row[i], COL_WIDTHS_DXA[COLUMNS[i]])

    return doc


def write_csv(leads: List[MatchedLead], out_path: str) -> str:
    """Machine-readable sidecar carrying every field, including the ones
    the .docx table has no room for.
    """
    with open(out_path, "w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=CSV_COLUMNS)
        writer.writeheader()
        for lead in leads:
            writer.writerow({
                "address": lead.address,
                "owner_name": lead.owner_name,
                "county": lead.county,
                "case_no": lead.case_numbers,
                "sources_hit": ";".join(lead.sources_hit),
                "distress_score": lead.distress_score,
                "address_source": lead.address_source,
                "address_confidence": lead.address_confidence or "",
                "hcad_matched_owner": lead.hcad_matched_owner or "",
                "parcel_id": lead.parcel_id or "",
                "mailing_address": lead.mailing_address or "",
                "market_value": lead.market_value or "",
                "year_built": lead.year_built or "",
                "building_sqft": lead.building_sqft or "",
                "source_urls": lead.source_urls,
            })
    return out_path


def write_docx(leads: List[MatchedLead], out_path: str, generated_at: datetime = None) -> str:
    doc = build_document(leads, generated_at)
    doc.save(out_path)
    return out_path
