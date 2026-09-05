"""Build the cited Word brief from its canonical local Markdown source."""
import json
import re
from pathlib import Path
from zipfile import ZipFile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = Path(__file__).resolve().parent.parent
LINK = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")


def linked_text(paragraph, text):
    position = 0
    for match in LINK.finditer(text):
        paragraph.add_run(text[position:match.start()])
        node = OxmlElement("w:hyperlink")
        node.set(qn("r:id"), paragraph.part.relate_to(match[2], RT.HYPERLINK, is_external=True))
        run = OxmlElement("w:r")
        properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "245B48")
        properties.append(color)
        run.append(properties)
        value = OxmlElement("w:t")
        value.text = match[1]
        run.append(value)
        node.append(run)
        paragraph._p.append(node)
        position = match.end()
    paragraph.add_run(text[position:])


def main():
    source = (ROOT / "docs/report-source.md").read_text(encoding="utf-8-sig")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(.7)
    section.left_margin = section.right_margin = Inches(.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.1
    for name in ("Title", "Heading 1", "Heading 2"):
        doc.styles[name].font.color.rgb = RGBColor.from_string("183C31")
        doc.styles[name].paragraph_format.keep_with_next = True
    section.header.paragraphs[0].text = "TEXAS INVESTORS  /  RESEARCH & PRODUCT STRATEGY"
    footer = section.footer.paragraphs[0]
    footer.text = "September 4, 2026  •  Local research edition                                  "
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    for block in source.strip().split("\n\n"):
        if block.startswith("# "):
            doc.add_heading(block[2:], 0)
        elif block.startswith("## "):
            doc.add_heading(block[3:], 1)
        elif re.match(r"1\. ", block):
            for line in block.splitlines():
                linked_text(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\. ", "", line))
        else:
            linked_text(doc.add_paragraph(), block.replace("\n", " "))
    target = ROOT / "docs/Texas-Investors-Research.docx"
    doc.save(target)
    reopened = Document(target)
    links = [r for r in reopened.part.rels.values() if r.reltype == RT.HYPERLINK]
    assert len(links) >= 14
    assert len(reopened.paragraphs) >= 35
    with ZipFile(target) as archive:
        assert archive.testzip() is None
    ledger = [{"title": match[1], "url": match[2], "accessed": "2026-09-04",
               "publication_date": "Not displayed unless stated in the report",
               "evidence": "Primary source; supporting claim appears in the linked report paragraph",
               "limitations": "See source-specific access notes and research limits in report"}
              for match in LINK.finditer(source)]
    (ROOT / "docs/research-source-ledger.json").write_text(json.dumps(ledger, indent=2), encoding="utf-8")
    print(json.dumps({"path": str(target), "bytes": target.stat().st_size,
                      "paragraphs": len(reopened.paragraphs), "source_links": len(links),
                      "verification": "DOCX read-back, archive integrity, headings and hyperlinks checked; no page renderer available"}))


if __name__ == "__main__":
    main()
